import os
import io
import fitz  # PyMuPDF
from PIL import Image
import google.generativeai as genai
from docx import Document

# === CONFIGURACIÓN ===
genai.configure(api_key="AIzaSyB3TQVzclIYarIK-Z0qDL0z6ERydQR-M-U")  # ⬅️ Reemplaza aquí con tu API key real

# === FUNCIONES ===

def convert_pdf_to_images(pdf_path, zoom=3.0):
    doc = fitz.open(pdf_path)
    images = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        images.append(img)
    return images

def extract_text_from_image(image):
    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content([
        "Extrae el texto de esta imagen y formátalo de manera clara y legible, respetando títulos, párrafos y listas.",
        image
    ])
    return response.text if response else ""

def ocr_pdf_to_docx(pdf_path):
    print(f"📄 Procesando OCR: {pdf_path}")
    images = convert_pdf_to_images(pdf_path)
    document = Document()

    for i, img in enumerate(images):
        print(f"🖼️ Página {i+1}/{len(images)} - Enviando a Gemini...")
        text = extract_text_from_image(img)
        document.add_paragraph(f"--- Página {i+1} ---", style="Heading 2")
        document.add_paragraph(text)
        document.add_page_break()

    output_path = pdf_path.replace(".pdf", "_OCR.docx")
    document.save(output_path)
    print(f"✅ OCR completado y guardado en: {output_path}")

# === USO ===
if __name__ == "__main__":
    # Ruta del archivo a procesar
    pdf_file = "Poliza_Chat/Poliza_Chatv1/Vida Jubilacion/Vida-Jubilacion.pdf"
    ocr_pdf_to_docx(pdf_file)
